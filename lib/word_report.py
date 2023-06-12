from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT

def generate_report(test_json,report_output_file_path):

    document = Document()

    document.add_heading('Test Notes', level=1)

    for test_row_json in test_json:
        print("Writing {}...".format(test_row_json.get('Number')))
        document.add_heading('{} - {}'.format(test_row_json.get('Number'),test_row_json.get('Requirement')), level=2)

        document.add_heading('Status', level=3)
        document.add_paragraph(test_row_json.get('Status'))

        document.add_heading('Tester: {} - {}'.format(test_row_json.get('Operator'),test_row_json.get('Date of Test')), level=3)

        document.add_heading('Reviewer:', level=3)

        document.add_heading('Comments', level=3)
        document.add_paragraph(test_row_json.get('method-comment'))

        document.add_heading('Criteria', level=3)
        document.add_paragraph(test_row_json.get('criteria-comment'))

        section = document.add_section(WD_SECTION.NEW_PAGE)

    document.save(report_output_file_path)

if __name__ =="__main__":

    from file_manager import *
   
    folder_path = 'data/123/'
    smo = "123"
    file_name = 'test_data.xlsx'

    all_tests_json = read_json_from_excel(folder_path,file_name)

    generate_report(test_json=all_tests_json,report_output_file_path="{folder_path}/testnotes_{smo}.docx".format(folder_path=folder_path,smo=smo))

